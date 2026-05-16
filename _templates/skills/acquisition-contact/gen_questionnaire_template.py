"""
Génère le document questionnaire vendeur pour le deal Talezzi.com
Format : socle fixe + section variable adaptée au deal
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Helpers ────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'), kwargs.get('val', 'single'))
        border.set(qn('w:sz'), str(kwargs.get('sz', 4)))
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), kwargs.get('color', 'CCCCCC'))
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_heading(doc, text, level=1, color_hex='1A3C5E'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14 if level == 1 else 11)
    r, g, b = int(color_hex[:2], 16), int(color_hex[2:4], 16), int(color_hex[4:], 16)
    run.font.color.rgb = RGBColor(r, g, b)
    if level == 1:
        # underline rule via paragraph border
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '1A3C5E')
        pBdr.append(bottom)
        pPr.append(pBdr)
    return p

def add_question_block(doc, number, question_text, response_lines=6):
    """Ajoute un bloc question + zone de réponse."""
    # Numéro + question
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run_num = p.add_run(f"Q{number}. ")
    run_num.bold = True
    run_num.font.size = Pt(10)
    run_num.font.color.rgb = RGBColor(0x1A, 0x3C, 0x5E)
    run_q = p.add_run(question_text)
    run_q.bold = True
    run_q.font.size = Pt(10)
    run_q.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    # Zone de réponse — tableau 1 cellule avec fond clair
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    set_cell_bg(cell, 'F5F8FB')
    set_cell_border(cell, val='single', sz=4, color='C5D5E8')
    cell.width = Inches(6.3)

    # Lignes de saisie (paragraphes vides)
    cell.paragraphs[0].paragraph_format.space_before = Pt(4)
    cell.paragraphs[0].paragraph_format.space_after = Pt(4)
    run_placeholder = cell.paragraphs[0].add_run("Votre réponse...")
    run_placeholder.font.size = Pt(10)
    run_placeholder.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
    run_placeholder.font.italic = True

    for _ in range(response_lines - 1):
        p2 = cell.add_paragraph("")
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(0)

    doc.add_paragraph()  # espacement


# ── Document ────────────────────────────────────────────────────────────────────

doc = Document()

# Marges
section = doc.sections[0]
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# Style par défaut
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(10)

# ── En-tête deal ───────────────────────────────────────────────────────────────

# Bandeau titre
title_table = doc.add_table(rows=1, cols=1)
title_cell = title_table.cell(0, 0)
set_cell_bg(title_cell, '1A3C5E')
title_cell.width = Inches(6.3)
tp = title_cell.paragraphs[0]
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tp.paragraph_format.space_before = Pt(12)
tp.paragraph_format.space_after = Pt(12)
tr = tp.add_run("QUESTIONNAIRE DE DUE DILIGENCE VENDEUR")
tr.bold = True
tr.font.size = Pt(14)
tr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
tr.font.name = 'Arial'

doc.add_paragraph()

# Fiche deal
info_table = doc.add_table(rows=4, cols=2)
info_table.style = 'Table Grid'
info_data = [
    ("Deal", "Talezzi.com — Plateforme ebooks enfants multilingue"),
    ("Source", "Flippa — listing #12767079"),
    ("Date", "16 mai 2026"),
    ("Statut", "🟡 À approfondir — Quick Screen validé"),
]
for i, (label, value) in enumerate(info_data):
    label_cell = info_table.cell(i, 0)
    value_cell = info_table.cell(i, 1)
    set_cell_bg(label_cell, 'E8EFF6')
    set_cell_border(label_cell, color='C5D5E8')
    set_cell_border(value_cell, color='C5D5E8')
    lp = label_cell.paragraphs[0]
    lp.paragraph_format.space_before = Pt(4)
    lp.paragraph_format.space_after = Pt(4)
    lr = lp.add_run(label)
    lr.bold = True
    lr.font.size = Pt(9)
    lr.font.name = 'Arial'
    vp = value_cell.paragraphs[0]
    vp.paragraph_format.space_before = Pt(4)
    vp.paragraph_format.space_after = Pt(4)
    vr = vp.add_run(value)
    vr.font.size = Pt(9)
    vr.font.name = 'Arial'

doc.add_paragraph()

# Note intro
intro = doc.add_paragraph()
intro.paragraph_format.space_before = Pt(4)
intro.paragraph_format.space_after = Pt(12)
intro_run = intro.add_run(
    "Ce document vous est partagé dans le cadre de notre processus de due diligence. "
    "Vos réponses nous aideront à avancer rapidement. Il n'y a pas de bonne ou mauvaise "
    "réponse — nous cherchons avant tout à comprendre le projet tel qu'il est."
)
intro_run.font.size = Pt(10)
intro_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
intro_run.font.italic = True

# ── SECTION A — Socle fixe ─────────────────────────────────────────────────────

add_heading(doc, "A — Questions générales (tous deals)", level=1)

add_question_block(doc, 1,
    "Pouvez-vous vous présenter brièvement ? Quel est votre parcours avec ce projet ?",
    response_lines=5)

add_question_block(doc, 2,
    "Quelle est la raison principale de la mise en vente aujourd'hui ?",
    response_lines=5)

add_question_block(doc, 3,
    "Tous les actifs listés dans l'annonce (code, domaine, contenus, comptes) sont-ils "
    "bien inclus dans la vente ? Y a-t-il des exclusions ?",
    response_lines=4)

add_question_block(doc, 4,
    "Existe-t-il des dettes, obligations contractuelles, litiges en cours ou engagements "
    "tiers liés à ce business dont nous devrions être informés ?",
    response_lines=4)

add_question_block(doc, 5,
    "Êtes-vous disponible pour accompagner la transition après la vente ? "
    "Sous quelle forme et sur quelle durée ?",
    response_lines=4)

# ── SECTION B — Questions spécifiques Talezzi ──────────────────────────────────

add_heading(doc, "B — Questions spécifiques à Talezzi.com", level=1)

note_b = doc.add_paragraph()
note_b.paragraph_format.space_after = Pt(8)
nb_run = note_b.add_run(
    "Ces questions portent sur les aspects qui nous semblent déterminants "
    "pour évaluer la valeur réelle de l'actif et la faisabilité de la reprise."
)
nb_run.font.size = Pt(10)
nb_run.font.italic = True
nb_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

add_heading(doc, "B1 — L'Universe Bible et le workflow de production", level=2, color_hex='2E6DA4')

add_question_block(doc, 6,
    "Pouvez-vous décrire concrètement la structure de l'Universe Bible ? "
    "Qu'est-ce qu'elle contient (personnages, univers, règles narratives, etc.) "
    "et comment est-elle organisée ?",
    response_lines=7)

add_question_block(doc, 7,
    "Quel est le workflow complet de production d'une nouvelle histoire, "
    "de l'idée initiale jusqu'à la publication sur la plateforme ? "
    "Combien de temps cela prend-il typiquement ?",
    response_lines=7)

add_question_block(doc, 8,
    "Les 3 histoires existantes et leurs traductions en 10 langues ont-elles été "
    "rédigées par vous, par des prestataires, ou générées (partiellement ou entièrement) "
    "par des outils IA ? Si IA, quel outil et dans quel cadre d'utilisation ?",
    response_lines=6)

add_heading(doc, "B2 — Droits de propriété intellectuelle", level=2, color_hex='2E6DA4')

add_question_block(doc, 9,
    "Disposez-vous de l'intégralité des droits sur les histoires, illustrations "
    "et autres contenus créatifs inclus dans la vente ? "
    "Ces droits sont-ils documentés par écrit (contrats de cession, factures, etc.) ?",
    response_lines=6)

add_question_block(doc, 10,
    "Le nom 'Talezzi', le logo et l'identité visuelle ont-ils fait l'objet "
    "d'un dépôt de marque ? Si non, savez-vous s'ils sont libres d'utilisation "
    "dans les marchés cibles (UE, UK) ?",
    response_lines=5)

add_heading(doc, "B3 — Technique et accès", level=2, color_hex='2E6DA4')

add_question_block(doc, 11,
    "Pouvez-vous nous donner un accès en lecture au dépôt GitHub privé "
    "dans le cadre de la due diligence ? Quelles sont vos conditions pour cela ?",
    response_lines=4)

add_question_block(doc, 12,
    "Stripe est intégré mais non activé. Y a-t-il une raison particulière "
    "(technique, légale, timing) ? L'activation est-elle prête à l'emploi ou "
    "nécessite-t-elle des développements supplémentaires ?",
    response_lines=5)

# ── Pied de page ───────────────────────────────────────────────────────────────

doc.add_paragraph()
footer_p = doc.add_paragraph()
footer_p.paragraph_format.space_before = Pt(16)
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Ligne de séparation
pPr = footer_p._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
top_border = OxmlElement('w:top')
top_border.set(qn('w:val'), 'single')
top_border.set(qn('w:sz'), '4')
top_border.set(qn('w:space'), '1')
top_border.set(qn('w:color'), 'C5D5E8')
pBdr.append(top_border)
pPr.append(pBdr)

fr = footer_p.add_run(
    "Document confidentiel — Acquisition Desk · InspirActions · 2026\n"
    "À compléter directement dans ce document et à retourner via Google Drive."
)
fr.font.size = Pt(8)
fr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
fr.font.italic = True

# ── Sauvegarde ─────────────────────────────────────────────────────────────────

output_path = "/sessions/relaxed-charming-davinci/mnt/ai-memory-test/notes/talezzi-questionnaire-vendeur.docx"
doc.save(output_path)
print(f"✅ Document généré : {output_path}")
